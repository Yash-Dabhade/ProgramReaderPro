# Text to speech
import os
import pyttsx3
import PyPDF2 as p
import docx

# initial setting of the engine for conversion

while True:
    def TtoS(str):
        engine.say(str)
        engine.runAndWait()
        engine.stop()
        engine.save_to_file(str, filename + '_REC.mp4')
        engine.runAndWait()

    def filetoopen(filename):
        with open(filename, 'r') as f:
            filecontent = f.read()
        TtoS(filecontent)

    def openword(filename):
        doc = docx.Document(filename)
        # all_paras = doc.paragraphs
        # # len(all_paras)
        # for para in all_paras:
        #     print(para.text)
        Text = ''
        for para in doc.paragraphs:
            Text += para.text
        with open('data.txt', 'w') as f:
            f.write(Text)
        filetoopen('data.txt')

    if __name__ == "__main__":
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
# selected voices[0] for men voice
        print("Please enter Speed for Talking between 20-200::")
        r = int(input())
        v = int(input("Enter 0 for male Voice or 1 for female voice::"))
        print(
            'Enter  file name(if in present folder) or enter full path(Without [" "]):')
        filename = input()
        engine.setProperty('voice', voices[v].id)
        engine.setProperty('rate', r)

        if v == 0:
            engine.say(
                "Hello  Friends. My name is Max and I am your FileReader !,Please Feel Free to talk with me. I was developed by Yash. Welcome !!! TO the Program Reader Pro .I will start reading in 3 seconds ,Please Wait")
            engine.runAndWait()
        if v == 1:
            engine.say(
                "Hello  Friends. My name is Lily and I am your FileReader !,Please Feel Free to talk with me. I was developed by Yash.Welcome !!! TO the Program Reader Pro . I will start reading in 3 seconds ,Please Wait")
            engine.runAndWait()

        if filename.endswith('pdf'):
            pdf = p.PdfFileReader(filename)
            s = ""
            for i in range(pdf.getNumPages()):
                s += pdf.getPage(i).extractText()
            with open('pdf1.txt', 'w', encoding='utf-8') as f:
                f.write(s)
            filetoopen('pdf1.txt')
            os.remove('pdf1.txt')
        elif filename.endswith('docx') or filename.endswith('doc'):
            openword(filename)
        else:
            filetoopen(filename)
        # os.remove('data.txt')
        input()
